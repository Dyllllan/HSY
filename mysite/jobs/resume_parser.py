"""
简历解析工具 - 提取PDF和DOCX文件中的文本内容
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_resume(file_path: str) -> str:
    """
    从简历文件中提取文本内容
    
    Args:
        file_path: 文件路径
    
    Returns:
        提取的文本内容
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.pdf':
            return _extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return _extract_from_docx(file_path)
        else:
            logger.warning(f"不支持的文件格式: {file_ext}")
            return ""
    except Exception as e:
        logger.error(f"提取简历文本失败: {str(e)}", exc_info=True)
        return ""


def _extract_from_pdf(file_path: str) -> str:
    """从PDF文件提取文本"""
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        return text.strip()
    except ImportError:
        logger.error("PyPDF2库未安装，请运行: pip install PyPDF2")
        # 尝试使用pdfplumber作为备选
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except ImportError:
            logger.error("pdfplumber库也未安装，请运行: pip install pdfplumber")
            return ""
    except Exception as e:
        logger.error(f"PDF解析失败: {str(e)}")
        return ""


def _extract_from_docx(file_path: str) -> str:
    """从DOCX文件提取文本"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except ImportError:
        logger.error("python-docx库未安装，请运行: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX解析失败: {str(e)}")
        return ""
